"""PKPM SATWE calculation report export — runtime.

Reads analysis results from PKPM SATWE output files (.OUT) and optionally
APIPyInterface.ResultData to produce a comprehensive calculation report
(JSON + Markdown + Word + PDF) covering design parameters, modal analysis,
earthquake forces, displacements, member design, and code checks.
"""
from __future__ import annotations

import os
import re
import subprocess
from pathlib import Path
from typing import Any, Dict, List, Optional


# ── Helpers ──────────────────────────────────────────────────────────────


def _safe_call(obj: Any, method: str, default: Any = None, *args: Any) -> Any:
    try:
        return getattr(obj, method)(*args)
    except Exception:
        return default


def _resolve_jws_path(model: Dict[str, Any], parameters: Dict[str, Any]) -> Path:
    jws = parameters.get("jws_path") or model.get("_pkpm_jws_path", "")
    if not jws:
        raise ValueError(
            "No JWS path provided. Pass parameters.jws_path or model._pkpm_jws_path."
        )
    p = Path(jws)
    if not p.is_file():
        raise FileNotFoundError(f"JWS file not found: {jws}")
    return p


# ── Project image discovery ─────────────────────────────────────────────


def _find_project_images(project_dir: Path) -> List[Path]:
    """Find structural images (BuildUp.BMP, PNG, JPG) in project directory."""
    images: List[Path] = []
    for pattern in ["BuildUp.BMP", "BuildUp.bmp", "*.png", "*.jpg", "*.jpeg"]:
        images.extend(project_dir.glob(pattern))
    # Deduplicate by stem, keep first found (BMP before PNG for same stem)
    seen: set[str] = set()
    result: List[Path] = []
    for p in images:
        key = p.stem.lower()
        if key not in seen:
            seen.add(key)
            result.append(p)
    return result


def _convert_bmp_to_png(bmp_path: Path) -> Path:
    """Convert BMP to PNG for embedding. Returns PNG path."""
    png_path = bmp_path.with_suffix(".png")
    if png_path.exists():
        return png_path
    try:
        from PIL import Image
        img = Image.open(str(bmp_path))
        if img.mode == "P":
            img = img.convert("RGB")
        img.save(str(png_path), "PNG")
        return png_path
    except Exception:
        return bmp_path


# ── .OUT file readers (GBK → structured text) ──────────────────────────


def _read_out_file(project_dir: Path, filename: str) -> Optional[str]:
    path = project_dir / filename
    if not path.is_file():
        return None
    try:
        raw = path.read_bytes()
        text = raw.decode("gbk", errors="replace")
        return text.replace("\r\n", "\n").replace("\r", "\n")
    except Exception:
        return None


def _read_all_wpj(project_dir: Path) -> List[Dict[str, Any]]:
    wpj_files = sorted(project_dir.glob("WPJ*.OUT"))
    results: List[Dict[str, Any]] = []
    for f in wpj_files:
        text = _read_out_file(project_dir, f.name)
        if text:
            results.append({"filename": f.name, "content": text})
    return results


def _parse_wmass_sections(text: str) -> Dict[str, str]:
    section_map: Dict[str, str] = {}
    patterns = [
        ("design_params", r"总信息\s+\.+\s*\n(.*?)(?=活荷载信息|二阶效应)"),
        ("wind_info_params", r"风荷载信息\s+\.+\s*\n(.*?)(?=地震信息)"),
        ("earthquake_params", r"地震信息\s+\.+\s*\n(.*?)(?=活荷载信息|二阶效应)"),
        ("live_load_params", r"活荷载信息\s+\.+\s*\n(.*?)(?=二阶效应|调整信息)"),
        ("second_order_params", r"二阶效应\s+\.+\s*\n(.*?)(?=调整信息)"),
        ("adjustment_params", r"调整信息\s+\.+\s*\n(.*?)(?=设计信息)"),
        ("design_info_params", r"设计信息\s+\.+\s*\n(.*?)(?=材料信息)"),
        ("material_params", r"材料信息\s+\.+\s*\n(.*?)(?=荷载组合信息)"),
        ("load_combination_params", r"荷载组合信息\s+\.+\s*\n(.*?)(?=地下信息)"),
        ("underground_params", r"地下信息\s+\.+\s*\n(.*?)(?=性能设计信息)"),
    ]
    for key, pat in patterns:
        m = re.search(pat, text, re.DOTALL)
        if m:
            section_map[key] = m.group(1).strip()
    return section_map


def _parse_mass_table(text: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    found_header = False
    for line in text.split("\n"):
        stripped = line.strip()
        if not found_header:
            if "\u5c42\u53f7" in stripped and "\u5854\u53f7" in stripped and "\u8d28\u5fc3" in stripped:
                found_header = True
            continue
        if not stripped:
            break
        parts = stripped.split()
        if len(parts) >= 8:
            try:
                int(parts[0])
                rows.append({
                    "floor": int(parts[0]),
                    "tower": int(parts[1]),
                    "mass_cx": float(parts[2]),
                    "mass_cy": float(parts[3]),
                    "mass_cz": float(parts[4]),
                    "dead_mass": float(parts[5]),
                    "live_mass": float(parts[6]),
                    "add_mass": float(parts[7]),
                    "mass_ratio": float(parts[8]) if len(parts) > 8 else 0.0,
                })
            except (ValueError, IndexError):
                break
    return rows


def _parse_total_mass(text: str) -> Dict[str, float]:
    result: Dict[str, float] = {}
    for key, pat in [
        ("live_total", r"活载产生的总质量\s*\(t\):\s*([\d.]+)"),
        ("dead_total", r"恒载产生的总质量\s*\(t\):\s*([\d.]+)"),
        ("add_total", r"附加总质量\s*\(t\):\s*([\d.]+)"),
        ("struct_total", r"结构的总质量\s*\(t\):\s*([\d.]+)"),
    ]:
        m = re.search(pat, text)
        if m:
            result[key] = float(m.group(1))
    return result


def _parse_member_counts(text: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    found_header = False
    for line in text.split("\n"):
        stripped = line.strip()
        if not found_header:
            if "\u5c42\u53f7" in stripped and "\u6881\u5143\u6570" in stripped:
                found_header = True
            continue
        if not stripped:
            if rows:
                break
            continue
        m = re.match(
            r"\s*(\d+)\(\s*(\d+)\)\s+(\d+)\s+(\d+)\(.*?\)\s+(\d+)\(.*?\)\s+(\d+)\(.*?\)\s+([\d.]+)\s+([\d.]+)",
            stripped,
        )
        if m:
            try:
                rows.append({
                    "floor": int(m.group(1)),
                    "std_floor": int(m.group(2)),
                    "tower": int(m.group(3)),
                    "beam_count": int(m.group(4)),
                    "column_count": int(m.group(5)),
                    "wall_count": int(m.group(6)),
                    "height": float(m.group(7)),
                    "cumulative_height": float(m.group(8)),
                })
            except (ValueError, IndexError):
                continue
    return rows


def _parse_wind_load(text: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    found_header = False
    for line in text.split("\n"):
        stripped = line.strip()
        if not found_header:
            if "\u98ce\u8377\u8f7dX" in stripped:
                found_header = True
            continue
        if not stripped:
            if rows:
                break
            continue
        if stripped.startswith("="):
            break
        parts = stripped.split()
        if len(parts) >= 8:
            try:
                int(parts[0])
                rows.append({
                    "floor": int(parts[0]),
                    "tower": int(parts[1]),
                    "wind_x": float(parts[2]),
                    "shear_x": float(parts[3]),
                    "overturn_x": float(parts[4]),
                    "wind_y": float(parts[5]),
                    "shear_y": float(parts[6]),
                    "overturn_y": float(parts[7]),
                })
            except (ValueError, IndexError):
                continue
    return rows


def _parse_floor_dimensions(text: str) -> List[Dict[str, Any]]:
    pattern = r"各楼层等效尺寸.*?\n\s*[-]+\s*\n(.*?)(?=\n\s*\*)"
    m = re.search(pattern, text, re.DOTALL)
    if not m:
        return []
    rows: List[Dict[str, Any]] = []
    for line in m.group(1).strip().split("\n"):
        parts = line.strip().split()
        if len(parts) >= 8:
            try:
                rows.append({
                    "floor": int(parts[0]),
                    "tower": int(parts[1]),
                    "area": float(parts[2]),
                    "centroid_x": float(parts[3]),
                    "centroid_y": float(parts[4]),
                    "equiv_width": float(parts[5]),
                    "equiv_height": float(parts[6]),
                    "max_width": float(parts[7]),
                    "min_width": float(parts[8]) if len(parts) > 8 else 0.0,
                })
            except (ValueError, IndexError):
                continue
    return rows


def _parse_unit_mass(text: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    found_data = False
    for line in text.split("\n"):
        stripped = line.strip()
        if not found_data:
            if "\u5355\u4f4d\u9762\u79ef\u8d28\u91cf" in stripped and "g[i]" in stripped:
                found_data = True
            continue
        if not stripped or stripped.startswith("="):
            break
        parts = stripped.split()
        if len(parts) >= 4:
            try:
                int(parts[0])
                rows.append({
                    "floor": int(parts[0]),
                    "tower": int(parts[1]),
                    "unit_mass": float(parts[2]),
                    "mass_ratio": float(parts[3]),
                })
            except (ValueError, IndexError):
                continue
    return rows


def _parse_stiffness_info(text: str) -> List[Dict[str, Any]]:
    pattern = r"Floor No\.\s+(\d+)\s+Tower No\.\s+(\d+)\s*\n(.*?)(?=-{20,}|\nX方向最小)"
    rows: List[Dict[str, Any]] = []
    for m in re.finditer(pattern, text, re.DOTALL):
        block = m.group(3)
        entry: Dict[str, Any] = {"floor": int(m.group(1)), "tower": int(m.group(2))}
        for key, pat in [
            ("Xstif", r"Xstif=\s*([\d.]+)"),
            ("Ystif", r"Ystif=\s*([\d.]+)"),
            ("Xmass", r"Xmass=\s*([\d.]+)"),
            ("Ymass", r"Ymass=\s*([\d.]+)"),
            ("Gmass", r"Gmass[^=]*=\s*([\d.]+)"),
            ("Eex", r"Eex\s*=\s*([\d.]+)"),
            ("Eey", r"Eey\s*=\s*([\d.]+)"),
            ("Ratx", r"Ratx\s*=\s*([\d.]+)"),
            ("Raty", r"Raty\s*=\s*([\d.]+)"),
            ("Ratx1", r"Ratx1=\s*([\d.]+)"),
            ("Raty1", r"Raty1=\s*([\d.]+)"),
            ("RJX1", r"RJX1\s*=\s*([\d.E+]+)"),
            ("RJY1", r"RJY1\s*=\s*([\d.E+]+)"),
            ("RJX3", r"RJX3\s*=\s*([\d.E+]+)"),
            ("RJY3", r"RJY3\s*=\s*([\d.E+]+)"),
        ]:
            m2 = re.search(pat, block)
            if m2:
                entry[key] = m2.group(1)
        rows.append(entry)
    return rows


def _parse_overturning(text: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    found_header = False
    for line in text.split("\n"):
        stripped = line.strip()
        if not found_header:
            if "\u6297\u503e\u8986\u529b\u77e9" in stripped:
                found_header = True
            continue
        if not stripped:
            if rows:
                break
            continue
        parts = stripped.split()
        if len(parts) >= 5:
            try:
                float(parts[2])
                rows.append({
                    "case": parts[0] + " " + parts[1],
                    "Mr": float(parts[2]),
                    "Mov": float(parts[3]),
                    "ratio": float(parts[4]),
                    "zero_stress": float(parts[5]) if len(parts) > 5 else 0.0,
                })
            except (ValueError, IndexError):
                continue
    return rows


def _parse_stability(text: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    found_header = False
    for line in text.split("\n"):
        stripped = line.strip()
        if not found_header:
            if "\u7ed3\u6784\u6574\u4f53\u7a33\u5b9a\u9a8c\u7b97\u7ed3\u679c" in stripped:
                found_header = True
            continue
        if "\u8be5\u7ed3\u6784\u521a\u91cd\u6bd4" in stripped:
            break
        if not stripped:
            continue
        parts = stripped.split()
        if len(parts) >= 7:
            try:
                int(parts[0])
                rows.append({
                    "floor": int(parts[0]),
                    "stiff_x": parts[1],
                    "stiff_y": parts[2],
                    "height": float(parts[3]),
                    "upper_weight": float(parts[4]),
                    "ratio_x": float(parts[5]),
                    "ratio_y": float(parts[6]),
                })
            except (ValueError, IndexError):
                continue
    return rows


def _parse_stability_conclusion(text: str) -> List[str]:
    results: List[str] = []
    for pat in [r"(该结构刚重比.*?验算)", r"(该结构刚重比.*?效应)"]:
        for m in re.finditer(pat, text):
            results.append(m.group(1))
    return results


def _parse_shear_capacity(text: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    found_header = False
    for line in text.split("\n"):
        stripped = line.strip()
        if not found_header:
            if "\u6297\u526a\u627f\u8f7d\u529b" in stripped and "\u627f\u8f7d\u529b\u6bd4\u503c" in stripped:
                found_header = True
            continue
        if not stripped:
            if rows:
                break
            continue
        if stripped.startswith("-"):
            continue
        parts = stripped.split()
        if len(parts) >= 6:
            try:
                int(parts[0])
                rows.append({
                    "floor": int(parts[0]),
                    "tower": int(parts[1]),
                    "capacity_x": parts[2],
                    "capacity_y": parts[3],
                    "ratio_x": parts[4],
                    "ratio_y": parts[5],
                })
            except (ValueError, IndexError):
                continue
    return rows


def _parse_comfort(text: str) -> List[str]:
    results: List[str] = []
    for line in text.split("\n"):
        if "加速度" in line and "=" in line:
            results.append(line.strip())
    return results


# ── WZQ.OUT parsers ─────────────────────────────────────────────────────


def _parse_wzq_periods(text: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    found_header = False
    for line in text.split("\n"):
        stripped = line.strip()
        if not found_header:
            if "\u632f\u578b\u53f7" in stripped and "\u5468 \u671f" in stripped:
                found_header = True
            continue
        if not stripped:
            if rows:
                break
            continue
        # Parse format: "  1       0.4524     90.00        1.00 ( 0.00+1.00 )      0.00"
        m = re.match(
            r"\s*(\d+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+\(\s*([-.\d]+)\+([-.\d]+)\s*\)\s+([-.\d]+)",
            stripped,
        )
        if m:
            try:
                rows.append({
                    "mode": int(m.group(1)),
                    "period": float(m.group(2)),
                    "angle": float(m.group(3)),
                    "translation": m.group(4),
                    "x_translation": m.group(5),
                    "y_translation": m.group(6),
                    "torsion": m.group(7),
                })
            except (ValueError, IndexError):
                continue
    return rows


def _parse_wzq_direction_factors(text: str) -> List[Dict[str, Any]]:
    pattern = r"直接输出X,Y,Z方向的.*?\n\s*[-]+\s*\n(.*?)(?=\n\s*=+|\n\s*\n)"
    m = re.search(pattern, text, re.DOTALL)
    if not m:
        return []
    rows: List[Dict[str, Any]] = []
    for line in m.group(1).strip().split("\n"):
        parts = line.strip().split()
        if len(parts) >= 5:
            try:
                int(parts[0])
                rows.append({
                    "mode": int(parts[0]),
                    "period": float(parts[1]),
                    "factor_x": float(parts[2]),
                    "factor_y": float(parts[3]),
                    "factor_z": float(parts[4]),
                })
            except (ValueError, IndexError):
                continue
    return rows


def _parse_wzq_base_shear(text: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    direction = None
    found_cqc_header = False
    past_note = False
    got_data = False
    for line in text.split("\n"):
        stripped = line.strip()
        if "(CQC)" in stripped and "\u697c\u5c42\u53cd\u5e94" not in stripped:
            if stripped.strip().startswith("X"):
                direction = "X"
            elif stripped.strip().startswith("Y"):
                direction = "Y"
            found_cqc_header = True
            past_note = False
            got_data = False
            continue
        if not found_cqc_header or not direction:
            continue
        if "\u6ce8\u610f" in stripped or "\u6ce8" in stripped:
            past_note = True
            continue
        if past_note and not got_data and not stripped:
            continue
        if past_note:
            if not stripped:
                if got_data:
                    found_cqc_header = False
                    direction = None
                continue
            if stripped.startswith("\u6297\u6807") or stripped.startswith("\u672c\u697c") or stripped.startswith("*") or stripped.startswith("="):
                found_cqc_header = False
                direction = None
                continue
            nums = re.findall(r"[-]?\d+\.\d+", stripped)
            ints_match = re.match(r"\s*(\d+)\s+(\d+)", stripped)
            if ints_match and len(nums) >= 2:
                try:
                    rows.append({
                        "direction": direction,
                        "floor": int(ints_match.group(1)),
                        "tower": int(ints_match.group(2)),
                        "F": float(nums[0]),
                        "V": float(nums[1]),
                    })
                    got_data = True
                except (ValueError, IndexError):
                    pass
    return rows


def _parse_wzq_effective_mass(text: str) -> Dict[str, List[Dict[str, Any]]]:
    result: Dict[str, List[Dict[str, Any]]] = {}
    for direction in ["X", "Y"]:
        pat = (
            rf"{direction}[^\n]*有效质量系数.*?\n\s*[-]+\s*\n(.*?)(?=\n\s*\n)"
        )
        m = re.search(pat, text, re.DOTALL)
        if m:
            entries: List[Dict[str, Any]] = []
            for line in m.group(1).strip().split("\n"):
                parts = line.strip().split()
                if len(parts) >= 2:
                    try:
                        entries.append({
                            "mode": int(parts[0]),
                            "coefficient": float(parts[1]),
                        })
                    except (ValueError, IndexError):
                        continue
            result[direction] = entries
    return result


def _parse_wzq_min_shear_ratio(text: str) -> Dict[str, str]:
    result: Dict[str, str] = {}
    for direction in ["X", "Y"]:
        pat = rf"抗标\(5\.2\.5\).*?{direction}向楼层最小剪重比\s*=\s*([\d.%]+)"
        m = re.search(pat, text)
        if m:
            result[direction] = m.group(1)
    return result


def _parse_wzq_shear_weight_ratio(text: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    found_header = False
    for line in text.split("\n"):
        stripped = line.strip()
        if not found_header:
            if "\u8c03\u6574\u7cfb\u6570" in stripped and "X" in stripped and "Y" in stripped:
                found_header = True
            continue
        if not stripped:
            if rows:
                break
            continue
        parts = stripped.split()
        if len(parts) >= 4:
            try:
                int(parts[0])
                rows.append({
                    "floor": int(parts[0]),
                    "tower": int(parts[1]),
                    "ratio_x": float(parts[2]),
                    "ratio_y": float(parts[3]),
                })
            except (ValueError, IndexError):
                continue
    return rows


# ── WDISP.OUT parsers ───────────────────────────────────────────────────


def _parse_wdisp_cases(text: str) -> List[Dict[str, Any]]:
    cases: List[Dict[str, Any]] = []
    for m in re.finditer(
        r"=== 工况\s+(\d+)\s+===\s*(.*?)(?=\n=== 工况|\Z)",
        text,
        re.DOTALL,
    ):
        case_title = m.group(2).split("\n")[0].strip()
        block = m.group(0)
        floor_data: List[Dict[str, Any]] = []
        for fm in re.finditer(
            r"(\d+)\s+(\d+)\s+(\d+)\s+([\d.-]+)\s+([\d.-]+)",
            block,
        ):
            try:
                floor_data.append({
                    "floor": int(fm.group(1)),
                    "tower": int(fm.group(2)),
                    "jmax": int(fm.group(3)),
                    "max_disp": float(fm.group(4)),
                    "ave_disp": float(fm.group(5)),
                })
            except (ValueError, IndexError):
                continue

        max_drift_line = ""
        dm = re.search(r"最大层间位移角:\s*(.*)", block)
        if dm:
            max_drift_line = dm.group(1).strip()

        cases.append({
            "case_num": int(m.group(1)),
            "title": case_title,
            "floors": floor_data,
            "max_drift_summary": max_drift_line,
        })
    return cases


# ── WGCPJ.OUT parser ────────────────────────────────────────────────────


def _parse_wgcpj(text: str) -> Dict[str, Any]:
    exceedance: List[str] = []
    for line in text.split("\n"):
        line = line.strip()
        if line and not line.startswith("-") and not line.startswith("|") and "SATWE" not in line and "2026" not in line and "WGCPJ" not in line and "工程项目" not in line and "项目编号" not in line and "第" in line:
            exceedance.append(line)
    return {"exceedance_count": len(exceedance), "items": exceedance}


# ── WPJ.OUT parser (per-floor member design) ────────────────────────────


def _parse_wpj_columns(text: str) -> List[Dict[str, Any]]:
    columns: List[Dict[str, Any]] = []
    for m in re.finditer(
        r"N-C=\s*(\d+)\s*\(\s*\d+\)\s*B\*H\(mm\)=\s*(\d+)\*\s*(\d+)(.*?)(?=N-C=|\Z)",
        text,
        re.DOTALL,
    ):
        block = m.group(4)
        entry: Dict[str, Any] = {
            "id": int(m.group(1)),
            "width": int(m.group(2)),
            "height": int(m.group(3)),
        }
        for key, pat in [
            ("cover", r"Cover=\s*(\d+)"),
            ("cx", r"Cx=\s*([\d.]+)"),
            ("cy", r"Cy=\s*([\d.]+)"),
            ("length", r"Lc=\s*([\d.]+)"),
            ("seismic_grade", r"Nfc=\s*(\d+)"),
            ("concrete_grade", r"Rcc=\s*([\d.]+)"),
            ("axial_ratio", r"Uc=\s*([\d.]+)"),
            ("reinforce_ratio", r"Rs=\s*([\d.]+)"),
            ("hoop_ratio", r"Rsv=\s*([\d.]+)"),
            ("corner_steel", r"Asc=\s*([\d.]+)"),
            ("shear_span", r"RMD=\s*([\d.]+)"),
        ]:
            m2 = re.search(pat, block)
            if m2:
                entry[key] = m2.group(1)
        columns.append(entry)
    return columns


def _parse_wpj_beams(text: str) -> List[Dict[str, Any]]:
    beams: List[Dict[str, Any]] = []
    for m in re.finditer(
        r"N-B=\s*(\d+)\s*\(.*?\)\s*B\*H\(mm\)=\s*(\d+)\*\s*(\d+)(.*?)(?=N-B=|\Z)",
        text,
        re.DOTALL,
    ):
        block = m.group(4)
        entry: Dict[str, Any] = {
            "id": int(m.group(1)),
            "width": int(m.group(2)),
            "height": int(m.group(3)),
        }
        for key, pat in [
            ("cover", r"Cover=\s*(\d+)"),
            ("length", r"Lb=\s*([\d.]+)"),
            ("seismic_grade", r"Nfb=\s*(\d+)"),
            ("concrete_grade", r"Rcb=\s*([\d.]+)"),
        ]:
            m2 = re.search(pat, block)
            if m2:
                entry[key] = m2.group(1)
        top_match = re.search(r"Top_Ast=\s*([\d.]+)", block)
        btm_match = re.search(r"Btm_Ast=\s*([\d.]+)", block)
        if top_match:
            entry["top_reinforce"] = top_match.group(1)
        if btm_match:
            entry["btm_reinforce"] = btm_match.group(1)
        beams.append(entry)
    return beams


# ── APIPyInterface extractors (supplementary) ───────────────────────────


def _extract_modal(result: Any) -> List[Dict[str, Any]]:
    periods = _safe_call(result, "GetModePeriods", [])
    out: List[Dict[str, Any]] = []
    for p in periods:
        out.append({
            "index": _safe_call(p, "GetIndex", 0),
            "period_s": round(_safe_call(p, "GetCycle", 0.0), 4),
            "angle": round(_safe_call(p, "GetAngle", 0.0), 2),
            "damping_ratio": round(_safe_call(p, "GetDampingRatio", 0.0), 4),
            "torsion_ratio": round(_safe_call(p, "GetTorsi", 0.0), 4),
            "x_side": round(_safe_call(p, "GetxSide", 0.0), 4),
            "y_side": round(_safe_call(p, "GetySide", 0.0), 4),
        })
    return out


def _extract_story_stiffness(result: Any) -> List[Dict[str, Any]]:
    data = _safe_call(result, "GetStoreyStifs", [])
    out: List[Dict[str, Any]] = []
    for s in data:
        out.append({
            "floor_index": _safe_call(s, "Getfloorindex", 0),
            "tower_index": _safe_call(s, "GetTowerIndex", 0),
            "RJX": round(_safe_call(s, "GetRJX", 0.0), 2),
            "RJY": round(_safe_call(s, "GetRJY", 0.0), 2),
            "ratio_x": round(_safe_call(s, "GetRatx", 0.0), 4),
            "ratio_y": round(_safe_call(s, "GetRaty", 0.0), 4),
        })
    return out


def _extract_story_drift(result: Any) -> Dict[str, Any]:
    eq_drift = _safe_call(result, "GetStoryDrift_Earthquake", {})
    wind_drift = _safe_call(result, "GetStoryDrift_Wind", {})
    drift_limit = _safe_call(result, "GetAngStoryDriftConsVal", 0.0)

    def _drift_entries(data: Any) -> Dict[str, List[Dict[str, Any]]]:
        if not isinstance(data, dict):
            return {}
        out: Dict[str, List[Dict[str, Any]]] = {}
        for key, entries in data.items():
            items: List[Dict[str, Any]] = []
            for e in entries:
                items.append({
                    "floor": _safe_call(e, "Getifloor", 0),
                    "tower": _safe_call(e, "Getitower", 0),
                    "loadcase": _safe_call(e, "Getiloadcase", 0),
                    "max_drift": round(_safe_call(e, "GetmaxD", 0.0), 4),
                    "min_drift": round(_safe_call(e, "GetminD", 0.0), 4),
                    "drift_angle": round(_safe_call(e, "GetratioA", 0.0), 6),
                    "inter_max_angle": round(_safe_call(e, "GetinterMaxA", 0.0), 6),
                })
            out[key] = items
        return out

    return {
        "earthquake": _drift_entries(eq_drift),
        "wind": _drift_entries(wind_drift),
        "limit_value": round(drift_limit, 4) if drift_limit else None,
    }


def _extract_base_shear(result: Any) -> Dict[str, Any]:
    data = _safe_call(result, "GetBearingShear", [])
    limit_val = _safe_call(result, "GetRatShearWeightConsVal", 0.0)
    out: List[Dict[str, Any]] = []
    for s in data:
        out.append({
            "floor": _safe_call(s, "GetFloorNum", 0),
            "tower": _safe_call(s, "GetTowerNum", 0),
            "ratio_x": round(_safe_call(s, "GetRatx", 0.0), 4),
            "ratio_y": round(_safe_call(s, "GetRaty", 0.0), 4),
            "limit_value": round(_safe_call(s, "GetLimitVal", 0.0), 4),
        })
    return {"entries": out, "shear_weight_limit": round(limit_val, 4) if limit_val else None}


def _extract_story_mass(result: Any) -> List[Dict[str, Any]]:
    data = _safe_call(result, "GetStoreyUnitMass", [])
    out: List[Dict[str, Any]] = []
    for s in data:
        out.append({
            "floor_index": _safe_call(s, "Getfloorindex", 0),
            "tower_index": _safe_call(s, "GetTowerIndex", 0),
            "unit_mass": round(_safe_call(s, "GetUnitMass", 0.0), 2),
            "area": round(_safe_call(s, "GetArea", 0.0), 2),
            "mass_ratio": round(_safe_call(s, "GetMassRatio", 0.0), 4),
            "poid_x": round(_safe_call(s, "GetpoidX", 0.0), 2),
            "poid_y": round(_safe_call(s, "GetpoidY", 0.0), 2),
        })
    return out


def _extract_stiff_weight_ratio(result: Any) -> Dict[str, Any]:
    data = _safe_call(result, "GetStiffWeightRatioFrame", [])
    limit_val = _safe_call(result, "GetRigidWeightRatioConsVal", 0.0)
    out: List[Dict[str, Any]] = []
    for s in data:
        out.append({
            "floor": _safe_call(s, "Getifloor", 0),
            "tower": _safe_call(s, "Getitower", 0),
            "height": round(_safe_call(s, "GetHeight", 0.0), 2),
            "stiff_x": round(_safe_call(s, "GetStiffX", 0.0), 2),
            "stiff_y": round(_safe_call(s, "GetStiffY", 0.0), 2),
            "ratio_x": round(_safe_call(s, "GetStiffWeightRatioX", 0.0), 4),
            "ratio_y": round(_safe_call(s, "GetStiffWeightRatioY", 0.0), 4),
        })
    return {"entries": out, "limit_value": round(limit_val, 4) if limit_val else None}


def _extract_beam_design(result: Any) -> Dict[str, Any]:
    beams_by_floor: Dict[int, List[Dict[str, Any]]] = {}
    max_shear_comp = 0.0
    total_reinforce = 0.0
    floor_idx = 1
    while True:
        beams = _safe_call(result, "GetDesignBeams", [], floor_idx)
        if not beams:
            break
        floor_beams: List[Dict[str, Any]] = []
        for b in beams:
            shear_ratio = _safe_call(b, "GetShearCompressionRatio", 0.0)
            max_shear_comp = max(max_shear_comp, shear_ratio)
            reinforce = _safe_call(b, "GetReinForceQuantity", 0.0)
            total_reinforce += reinforce
            floor_beams.append({
                "pmid": _safe_call(b, "GetPmid", 0),
                "shear_compression_ratio": round(shear_ratio, 4),
                "reinforce_quantity": round(reinforce, 2),
                "concrete_quantity": round(_safe_call(b, "GetConcreteQuantity", 0.0), 2),
                "steel_quantity": round(_safe_call(b, "GetSteelQuantity", 0.0), 2),
            })
        beams_by_floor[floor_idx] = floor_beams
        floor_idx += 1
    return {
        "floors": beams_by_floor,
        "total_beams": sum(len(v) for v in beams_by_floor.values()),
        "max_shear_compression_ratio": round(max_shear_comp, 4),
        "total_reinforce_quantity": round(total_reinforce, 2),
        "floors_analyzed": floor_idx - 1,
    }


def _extract_column_design(result: Any) -> Dict[str, Any]:
    columns_by_floor: Dict[int, List[Dict[str, Any]]] = {}
    max_axial = 0.0
    total_reinforce = 0.0
    floor_idx = 1
    while True:
        columns = _safe_call(result, "GetDesignColumns", [], floor_idx)
        if not columns:
            break
        floor_columns: List[Dict[str, Any]] = []
        for c in columns:
            axial_ratios = _safe_call(c, "GetAxialCompresRatio", [0.0])
            axial_max = max(axial_ratios) if axial_ratios else 0.0
            max_axial = max(max_axial, axial_max)
            reinforce = _safe_call(c, "GetReinForceQuantity", 0.0)
            total_reinforce += reinforce
            floor_columns.append({
                "pmid": _safe_call(c, "GetPmid", 0),
                "element_id": _safe_call(c, "GetElementid", 0),
                "axial_compression_ratio": [round(v, 4) for v in axial_ratios],
                "max_axial_compression": round(axial_max, 4),
                "reinforce_quantity": round(reinforce, 2),
                "concrete_quantity": round(_safe_call(c, "GetConcreteQuantity", 0.0), 2),
                "steel_quantity": round(_safe_call(c, "GetSteelQuantity", 0.0), 2),
                "slender_ratio": [round(v, 4) for v in _safe_call(c, "GetSlenderRatio", [0.0])],
            })
        columns_by_floor[floor_idx] = floor_columns
        floor_idx += 1
    return {
        "floors": columns_by_floor,
        "total_columns": sum(len(v) for v in columns_by_floor.values()),
        "max_axial_compression_ratio": round(max_axial, 4),
        "total_reinforce_quantity": round(total_reinforce, 2),
        "floors_analyzed": floor_idx - 1,
    }


# ── Markdown generation ─────────────────────────────────────────────────


def _kv_lines(text: str) -> List[str]:
    lines: List[str] = []
    for line in text.strip().split("\n"):
        line = line.strip()
        if line:
            lines.append(f"- {line}")
    return lines


def _generate_markdown(report: Dict[str, Any]) -> str:
    lines: List[str] = []
    lines.append("# PKPM SATWE 结构计算书")
    lines.append("")

    # Structural images
    images: List[str] = report.get("detailed", {}).get("images", [])
    if images:
        lines.append("## 结构平面图")
        lines.append("")
        for img in images:
            lines.append(f"![结构平面图]({img})")
            lines.append("")

    detailed = report.get("detailed", {})
    out_data = detailed.get("out_file_data", {})

    # Section 1: Design Parameters
    params = out_data.get("wmass_params", {})
    if params:
        lines.append("## 一、设计参数总信息")
        lines.append("")
        for key, label in [
            ("design_params", "总体信息"),
            ("wind_info_params", "风荷载参数"),
            ("earthquake_params", "地震参数"),
            ("material_params", "材料参数"),
            ("adjustment_params", "调整参数"),
            ("design_info_params", "设计信息"),
            ("load_combination_params", "荷载组合参数"),
        ]:
            section = params.get(key)
            if section:
                lines.append(f"### {label}")
                lines.append("")
                lines.extend(_kv_lines(section))
                lines.append("")

    # Section 2: Mass & Member Counts
    mass_table = out_data.get("mass_table", [])
    member_counts = out_data.get("member_counts", [])
    total_mass = out_data.get("total_mass", {})
    if mass_table or member_counts:
        lines.append("## 二、楼层质量与构件信息")
        lines.append("")
        if total_mass:
            for k, v in total_mass.items():
                label = {
                    "live_total": "活载总质量",
                    "dead_total": "恒载总质量",
                    "add_total": "附加总质量",
                    "struct_total": "结构总质量",
                }.get(k, k)
                lines.append(f"- **{label}**: {v} t")
            lines.append("")
        if mass_table:
            lines.append("### 各层质量")
            lines.append("")
            lines.append("| 层号 | 塔号 | 恒载质量(t) | 活载质量(t) | 质量比 |")
            lines.append("|------|------|------------|------------|--------|")
            for r in mass_table:
                lines.append(f"| {r['floor']} | {r['tower']} | {r['dead_mass']} | {r['live_mass']} | {r['mass_ratio']} |")
            lines.append("")
        if member_counts:
            lines.append("### 构件数量与层高")
            lines.append("")
            lines.append("| 层号 | 梁数 | 柱数 | 层高(m) | 累计高度(m) |")
            lines.append("|------|------|------|---------|------------|")
            for r in member_counts:
                lines.append(f"| {r['floor']} | {r['beam_count']} | {r['column_count']} | {r['height']} | {r['cumulative_height']} |")
            lines.append("")

    # Section 3: Unit Mass
    unit_mass = out_data.get("unit_mass", [])
    if unit_mass:
        lines.append("## 三、单位面积质量分布")
        lines.append("")
        lines.append("| 层号 | 塔号 | 单位面积质量(kg/m²) | 质量比 |")
        lines.append("|------|------|--------------------|--------|")
        for r in unit_mass:
            lines.append(f"| {r['floor']} | {r['tower']} | {r['unit_mass']} | {r['mass_ratio']} |")
        lines.append("")

    # Section 4: Wind Load
    wind_load = out_data.get("wind_load", [])
    if wind_load:
        lines.append("## 四、风荷载信息")
        lines.append("")
        lines.append("| 层号 | 风荷载X(kN) | 剪力X(kN) | 倾覆X(kN·m) | 风荷载Y(kN) | 剪力Y(kN) | 倾覆Y(kN·m) |")
        lines.append("|------|------------|----------|------------|------------|----------|------------|")
        for r in wind_load:
            lines.append(
                f"| {r['floor']} | {r['wind_x']} | {r['shear_x']} | {r['overturn_x']} | "
                f"{r['wind_y']} | {r['shear_y']} | {r['overturn_y']} |"
            )
        lines.append("")

    # Section 5: Modal Analysis
    modal_api = detailed.get("modal_analysis", [])
    wzq_periods = out_data.get("wzq_periods", [])
    modal_data = wzq_periods if wzq_periods else modal_api
    if modal_data:
        lines.append("## 五、模态分析")
        lines.append("")
        if wzq_periods:
            lines.append("| 振型号 | 周期(s) | 方向角(°) | 平动系数 | X平动 | Y平动 | 扭转系数 |")
            lines.append("|--------|---------|-----------|---------|-------|-------|---------|")
            for m in wzq_periods:
                lines.append(
                    f"| {m['mode']} | {m['period']} | {m['angle']} | {m['translation']} | "
                    f"{m['x_translation']} | {m['y_translation']} | {m['torsion']} |"
                )
        else:
            lines.append("| 振型号 | 周期(s) | 方向角(°) | 扭转系数 | X侧移 | Y侧移 |")
            lines.append("|--------|---------|-----------|----------|-------|-------|")
            for m in modal_api:
                lines.append(
                    f"| {m['index']} | {m['period_s']} | {m['angle']} | "
                    f"{m['torsion_ratio']} | {m['x_side']} | {m['y_side']} |"
                )
        lines.append("")

    # Section 6: Earthquake Response (CQC)
    wzq_base_shear = out_data.get("wzq_base_shear", [])
    if wzq_base_shear:
        lines.append("## 六、地震作用下的楼层反应（CQC）")
        lines.append("")
        for direction in ["X", "Y"]:
            dir_data = [r for r in wzq_base_shear if r["direction"] == direction]
            if dir_data:
                lines.append(f"### {direction}方向")
                lines.append("")
                lines.append("| 层号 | 塔号 | F(kN) | V(kN) |")
                lines.append("|------|------|-------|-------|")
                for r in dir_data:
                    lines.append(f"| {r['floor']} | {r['tower']} | {r['F']} | {r['V']} |")
                lines.append("")

    # Section 7: Shear Weight Ratio
    shear_weight = out_data.get("wzq_shear_weight_ratio", [])
    if shear_weight:
        lines.append("## 七、楼层剪重比")
        lines.append("")
        lines.append("| 层号 | 塔号 | X向剪重比 | Y向剪重比 |")
        lines.append("|------|------|----------|----------|")
        for r in shear_weight:
            lines.append(f"| {r['floor']} | {r['tower']} | {r['ratio_x']} | {r['ratio_y']} |")
        lines.append("")

    # Section 8: Floor Displacement
    wdisp_cases = out_data.get("wdisp_cases", [])
    if wdisp_cases:
        lines.append("## 八、层间位移")
        lines.append("")
        for case in wdisp_cases:
            lines.append(f"### 工况 {case['case_num']}: {case['title']}")
            lines.append("")
            if case["floors"]:
                lines.append("| 层号 | 塔号 | 最大位移(mm) | 平均位移(mm) |")
                lines.append("|------|------|------------|------------|")
                for f in case["floors"]:
                    lines.append(f"| {f['floor']} | {f['tower']} | {f['max_disp']} | {f['ave_disp']} |")
            if case["max_drift_summary"]:
                lines.append(f"\n**最大层间位移角**: {case['max_drift_summary']}")
            lines.append("")

    # Section 9: Story Stiffness
    stiff_api = detailed.get("story_stiffness", [])
    stiffness_info = out_data.get("stiffness_info", [])
    if stiff_api or stiffness_info:
        lines.append("## 九、层刚度")
        lines.append("")
        if stiff_api:
            lines.append("| 层号 | RJX(kN/m) | RJY(kN/m) | 比值X | 比值Y |")
            lines.append("|------|-----------|-----------|-------|-------|")
            for s in stiff_api:
                lines.append(
                    f"| {s['floor_index']} | {s['RJX']} | {s['RJY']} | "
                    f"{s['ratio_x']} | {s['ratio_y']} |"
                )
        if stiffness_info:
            lines.append("\n### 刚心、偏心率信息")
            lines.append("")
            lines.append("| 层号 | 刚心X | 刚心Y | 质心X | 质心Y | 偏心率X | 偏心率Y |")
            lines.append("|------|-------|-------|-------|-------|--------|--------|")
            for s in stiffness_info:
                lines.append(
                    f"| {s['floor']} | {s.get('Xstif', '')} | {s.get('Ystif', '')} | "
                    f"{s.get('Xmass', '')} | {s.get('Ymass', '')} | "
                    f"{s.get('Eex', '')} | {s.get('Eey', '')} |"
                )
        lines.append("")

    # Section 10: Overturning
    overturning = out_data.get("overturning", [])
    if overturning:
        lines.append("## 十、抗倾覆验算")
        lines.append("")
        lines.append("| 工况 | 抗倾覆力矩 | 倾覆力矩 | 比值 | 零应力区(%) |")
        lines.append("|------|-----------|---------|------|------------|")
        for r in overturning:
            lines.append(f"| {r['case']} | {r['Mr']} | {r['Mov']} | {r['ratio']} | {r['zero_stress']} |")
        lines.append("")

    # Section 11: Stability
    stability = out_data.get("stability", [])
    stability_conclusion = out_data.get("stability_conclusion", [])
    if stability:
        lines.append("## 十一、结构整体稳定验算")
        lines.append("")
        lines.append("| 层号 | X刚度 | Y刚度 | 层高 | 上部重量 | X刚重比 | Y刚重比 |")
        lines.append("|------|-------|-------|------|---------|--------|--------|")
        for r in stability:
            lines.append(
                f"| {r['floor']} | {r['stiff_x']} | {r['stiff_y']} | {r['height']} | "
                f"{r['upper_weight']} | {r['ratio_x']} | {r['ratio_y']} |"
            )
        lines.append("")
        for c in stability_conclusion:
            lines.append(f"- {c}")
        lines.append("")

    # Section 12: Comfort
    comfort = out_data.get("comfort", [])
    if comfort:
        lines.append("## 十二、结构舒适性验算")
        lines.append("")
        for c in comfort:
            lines.append(f"- {c}")
        lines.append("")

    # Section 13: Shear Capacity
    shear_cap = out_data.get("shear_capacity", [])
    if shear_cap:
        lines.append("## 十三、楼层抗剪承载力")
        lines.append("")
        lines.append("| 层号 | 塔号 | X向承载力 | Y向承载力 | 比值X | 比值Y |")
        lines.append("|------|------|----------|----------|-------|-------|")
        for r in shear_cap:
            lines.append(
                f"| {r['floor']} | {r['tower']} | {r['capacity_x']} | "
                f"{r['capacity_y']} | {r['ratio_x']} | {r['ratio_y']} |"
            )
        lines.append("")

    # Section 14: Column Design (from WPJ)
    wpj_columns = out_data.get("wpj_columns", [])
    if wpj_columns:
        lines.append("## 十四、柱配筋验算结果")
        lines.append("")
        lines.append("| 编号 | 截面B*H(mm) | 轴压比 | 配筋率(%) | 配箍率(%) | 角筋面积 |")
        lines.append("|------|------------|--------|----------|----------|---------|")
        for c in wpj_columns:
            lines.append(
                f"| {c['id']} | {c.get('width', '')}*{c.get('height', '')} | "
                f"{c.get('axial_ratio', '')} | {c.get('reinforce_ratio', '')} | "
                f"{c.get('hoop_ratio', '')} | {c.get('corner_steel', '')} |"
            )
        lines.append("")

    # Section 15: Beam Design (from WPJ)
    wpj_beams = out_data.get("wpj_beams", [])
    if wpj_beams:
        lines.append("## 十五、梁配筋验算结果")
        lines.append("")
        lines.append("| 编号 | 截面B*H(mm) | 上部配筋 | 下部配筋 |")
        lines.append("|------|------------|---------|---------|")
        for b in wpj_beams:
            lines.append(
                f"| {b['id']} | {b.get('width', '')}*{b.get('height', '')} | "
                f"{b.get('top_reinforce', '')} | {b.get('btm_reinforce', '')} |"
            )
        lines.append("")

    # Section 16: Exceedance
    wgcpj = out_data.get("wgcpj", {})
    exceed_items = wgcpj.get("items", [])
    if exceed_items:
        lines.append("## 十六、超限信息")
        lines.append("")
        for item in exceed_items:
            lines.append(f"- {item}")
        lines.append("")

    # Section 17: Beam/Column Summary (API data)
    beams = detailed.get("beam_design", {})
    if beams.get("total_beams", 0) > 0:
        lines.append("## 十七、梁设计统计")
        lines.append("")
        lines.append(f"- **总梁数**: {beams['total_beams']}")
        lines.append(f"- **最大剪压比**: {beams['max_shear_compression_ratio']}")
        lines.append(f"- **总配筋量**: {beams['total_reinforce_quantity']}")
        lines.append("")

    cols = detailed.get("column_design", {})
    if cols.get("total_columns", 0) > 0:
        lines.append("## 十八、柱设计统计")
        lines.append("")
        lines.append(f"- **总柱数**: {cols['total_columns']}")
        lines.append(f"- **最大轴压比**: {cols['max_axial_compression_ratio']}")
        lines.append(f"- **总配筋量**: {cols['total_reinforce_quantity']}")
        lines.append("")

    return "\n".join(lines)


# ── Word document generation ────────────────────────────────────────────


def _add_table(doc: Any, headers: List[str], rows: List[List[str]]) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = 1
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "D9E2F3",
            qn("w:val"): "clear",
        })
        shading.append(shading_elem)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = 1
                for run in p.runs:
                    run.font.size = Pt(10)


def _add_kv_section(doc: Any, title: str, text: str) -> None:
    doc.add_heading(title, level=2)
    for line in text.strip().split("\n"):
        line = line.strip()
        if line:
            doc.add_paragraph(line, style="List Bullet")


def _generate_docx(report: Dict[str, Any], output_path: Path, images: Optional[List[Path]] = None) -> Path:
    from docx.shared import Inches, Pt
    from docx import Document

    doc = Document()
    title = doc.add_heading("PKPM SATWE 结构计算书", level=0)
    title.alignment = 1
    for run in title.runs:
        run.font.size = Pt(22)

    # Insert structural images after title
    if images:
        doc.add_heading("结构平面图", level=1)
        for img_path in images:
            embed_path = _convert_bmp_to_png(img_path) if img_path.suffix.lower() == ".bmp" else img_path
            try:
                doc.add_picture(str(embed_path), width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = 1
            except Exception:
                pass

    detailed = report.get("detailed", {})
    out_data = detailed.get("out_file_data", {})
    params = out_data.get("wmass_params", {})

    # Section 1
    if params:
        doc.add_heading("一、设计参数总信息", level=1)
        for key, label in [
            ("design_params", "总体信息"),
            ("wind_info_params", "风荷载参数"),
            ("earthquake_params", "地震参数"),
            ("material_params", "材料参数"),
        ]:
            section = params.get(key)
            if section:
                _add_kv_section(doc, label, section)

    # Section 2: Mass
    mass_table = out_data.get("mass_table", [])
    total_mass = out_data.get("total_mass", {})
    member_counts = out_data.get("member_counts", [])
    if mass_table or member_counts:
        doc.add_heading("二、楼层质量与构件信息", level=1)
        if total_mass:
            for k, v in total_mass.items():
                doc.add_paragraph(f"{k}: {v} t")
        if mass_table:
            doc.add_heading("各层质量", level=2)
            _add_table(doc,
                ["层号", "塔号", "恒载质量(t)", "活载质量(t)", "质量比"],
                [[str(r["floor"]), str(r["tower"]), str(r["dead_mass"]),
                  str(r["live_mass"]), str(r["mass_ratio"])] for r in mass_table])
        if member_counts:
            doc.add_heading("构件数量与层高", level=2)
            _add_table(doc,
                ["层号", "梁数", "柱数", "层高(m)", "累计高度(m)"],
                [[str(r["floor"]), str(r["beam_count"]), str(r["column_count"]),
                  str(r["height"]), str(r["cumulative_height"])] for r in member_counts])

    # Section 3: Unit Mass
    unit_mass = out_data.get("unit_mass", [])
    if unit_mass:
        doc.add_heading("三、单位面积质量分布", level=1)
        _add_table(doc,
            ["层号", "塔号", "单位面积质量(kg/m²)", "质量比"],
            [[str(r["floor"]), str(r["tower"]), str(r["unit_mass"]), str(r["mass_ratio"])]
             for r in unit_mass])

    # Section 4: Wind Load
    wind_load = out_data.get("wind_load", [])
    if wind_load:
        doc.add_heading("四、风荷载信息", level=1)
        _add_table(doc,
            ["层号", "风荷载X", "剪力X", "倾覆X", "风荷载Y", "剪力Y", "倾覆Y"],
            [[str(r["floor"]), str(r["wind_x"]), str(r["shear_x"]), str(r["overturn_x"]),
              str(r["wind_y"]), str(r["shear_y"]), str(r["overturn_y"])] for r in wind_load])

    # Section 5: Modal
    wzq_periods = out_data.get("wzq_periods", [])
    modal_api = detailed.get("modal_analysis", [])
    if wzq_periods:
        doc.add_heading("五、模态分析", level=1)
        _add_table(doc,
            ["振型号", "周期(s)", "方向角(°)", "平动系数", "X平动", "Y平动", "扭转系数"],
            [[str(m["mode"]), str(m["period"]), str(m["angle"]), str(m["translation"]),
              str(m["x_translation"]), str(m["y_translation"]), str(m["torsion"])]
             for m in wzq_periods])
    elif modal_api:
        doc.add_heading("五、模态分析", level=1)
        _add_table(doc,
            ["振型号", "周期(s)", "方向角(°)", "扭转系数", "X侧移", "Y侧移"],
            [[str(m["index"]), str(m["period_s"]), str(m["angle"]),
              str(m["torsion_ratio"]), str(m["x_side"]), str(m["y_side"])]
             for m in modal_api])

    # Section 6: Earthquake CQC
    wzq_base_shear = out_data.get("wzq_base_shear", [])
    if wzq_base_shear:
        doc.add_heading("六、地震作用下的楼层反应（CQC）", level=1)
        for direction in ["X", "Y"]:
            dir_data = [r for r in wzq_base_shear if r["direction"] == direction]
            if dir_data:
                doc.add_heading(f"{direction}方向", level=2)
                _add_table(doc,
                    ["层号", "塔号", "F(kN)", "V(kN)"],
                    [[str(r["floor"]), str(r["tower"]), str(r["F"]), str(r["V"])]
                     for r in dir_data])

    # Section 7: Shear Weight Ratio
    shear_weight = out_data.get("wzq_shear_weight_ratio", [])
    if shear_weight:
        doc.add_heading("七、楼层剪重比", level=1)
        _add_table(doc,
            ["层号", "塔号", "X向剪重比", "Y向剪重比"],
            [[str(r["floor"]), str(r["tower"]), str(r["ratio_x"]), str(r["ratio_y"])]
             for r in shear_weight])

    # Section 8: Displacement
    wdisp_cases = out_data.get("wdisp_cases", [])
    if wdisp_cases:
        doc.add_heading("八、层间位移", level=1)
        for case in wdisp_cases:
            doc.add_heading(f"工况 {case['case_num']}: {case['title']}", level=2)
            if case["floors"]:
                _add_table(doc,
                    ["层号", "塔号", "最大位移(mm)", "平均位移(mm)"],
                    [[str(f["floor"]), str(f["tower"]), str(f["max_disp"]), str(f["ave_disp"])]
                     for f in case["floors"]])
            if case["max_drift_summary"]:
                doc.add_paragraph(f"最大层间位移角: {case['max_drift_summary']}")

    # Section 9: Stiffness
    stiff_api = detailed.get("story_stiffness", [])
    if stiff_api:
        doc.add_heading("九、层刚度", level=1)
        _add_table(doc,
            ["层号", "RJX(kN/m)", "RJY(kN/m)", "比值X", "比值Y"],
            [[str(s["floor_index"]), str(s["RJX"]), str(s["RJY"]),
              str(s["ratio_x"]), str(s["ratio_y"])] for s in stiff_api])

    # Section 10: Overturning
    overturning = out_data.get("overturning", [])
    if overturning:
        doc.add_heading("十、抗倾覆验算", level=1)
        _add_table(doc,
            ["工况", "抗倾覆力矩", "倾覆力矩", "比值", "零应力区(%)"],
            [[str(r["case"]), str(r["Mr"]), str(r["Mov"]), str(r["ratio"]), str(r["zero_stress"])]
             for r in overturning])

    # Section 11: Stability
    stability = out_data.get("stability", [])
    if stability:
        doc.add_heading("十一、结构整体稳定验算", level=1)
        _add_table(doc,
            ["层号", "X刚度", "Y刚度", "层高", "上部重量", "X刚重比", "Y刚重比"],
            [[str(r["floor"]), str(r["stiff_x"]), str(r["stiff_y"]), str(r["height"]),
              str(r["upper_weight"]), str(r["ratio_x"]), str(r["ratio_y"])]
             for r in stability])
        for c in out_data.get("stability_conclusion", []):
            doc.add_paragraph(c)

    # Section 12: Comfort
    comfort = out_data.get("comfort", [])
    if comfort:
        doc.add_heading("十二、结构舒适性验算", level=1)
        for c in comfort:
            doc.add_paragraph(c, style="List Bullet")

    # Section 13: Shear Capacity
    shear_cap = out_data.get("shear_capacity", [])
    if shear_cap:
        doc.add_heading("十三、楼层抗剪承载力", level=1)
        _add_table(doc,
            ["层号", "塔号", "X向承载力", "Y向承载力", "比值X", "比值Y"],
            [[str(r["floor"]), str(r["tower"]), str(r["capacity_x"]),
              str(r["capacity_y"]), str(r["ratio_x"]), str(r["ratio_y"])]
             for r in shear_cap])

    # Section 14: Column Design
    wpj_columns = out_data.get("wpj_columns", [])
    if wpj_columns:
        doc.add_heading("十四、柱配筋验算结果", level=1)
        _add_table(doc,
            ["编号", "截面B*H(mm)", "轴压比", "配筋率(%)", "配箍率(%)", "角筋面积"],
            [[str(c["id"]), f"{c.get('width', '')}*{c.get('height', '')}",
              str(c.get("axial_ratio", "")), str(c.get("reinforce_ratio", "")),
              str(c.get("hoop_ratio", "")), str(c.get("corner_steel", ""))]
             for c in wpj_columns])

    # Section 15: Beam Design
    wpj_beams = out_data.get("wpj_beams", [])
    if wpj_beams:
        doc.add_heading("十五、梁配筋验算结果", level=1)
        _add_table(doc,
            ["编号", "截面B*H(mm)", "上部配筋", "下部配筋"],
            [[str(b["id"]), f"{b.get('width', '')}*{b.get('height', '')}",
              str(b.get("top_reinforce", "")), str(b.get("btm_reinforce", ""))]
             for b in wpj_beams])

    # Section 16: Exceedance
    wgcpj = out_data.get("wgcpj", {})
    exceed_items = wgcpj.get("items", [])
    if exceed_items:
        doc.add_heading("十六、超限信息", level=1)
        for item in exceed_items:
            doc.add_paragraph(item, style="List Bullet")

    # Section 17-18: API Summary
    beams = detailed.get("beam_design", {})
    if beams.get("total_beams", 0) > 0:
        doc.add_heading("十七、梁设计统计", level=1)
        doc.add_paragraph(f"总梁数: {beams['total_beams']}")
        doc.add_paragraph(f"最大剪压比: {beams['max_shear_compression_ratio']}")
        doc.add_paragraph(f"总配筋量: {beams['total_reinforce_quantity']}")

    cols = detailed.get("column_design", {})
    if cols.get("total_columns", 0) > 0:
        doc.add_heading("十八、柱设计统计", level=1)
        doc.add_paragraph(f"总柱数: {cols['total_columns']}")
        doc.add_paragraph(f"最大轴压比: {cols['max_axial_compression_ratio']}")
        doc.add_paragraph(f"总配筋量: {cols['total_reinforce_quantity']}")

    doc.save(str(output_path))
    return output_path


# ── PDF generation ──────────────────────────────────────────────────────


def _generate_pdf(report: Dict[str, Any], output_path: Path, images: Optional[List[Path]] = None) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    cn_font = "STSong-Light"

    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CNTitle", parent=styles["Title"], fontName=cn_font, fontSize=18, spaceAfter=12)
    h1_style = ParagraphStyle("CNH1", parent=styles["Heading1"], fontName=cn_font, fontSize=14, spaceAfter=8, spaceBefore=12)
    h2_style = ParagraphStyle("CNH2", parent=styles["Heading2"], fontName=cn_font, fontSize=12, spaceAfter=6, spaceBefore=8)
    body_style = ParagraphStyle("CNBody", parent=styles["Normal"], fontName=cn_font, fontSize=10, spaceAfter=4)

    elements: List[Any] = []
    detailed = report.get("detailed", {})
    out_data = detailed.get("out_file_data", {})

    elements.append(Paragraph("PKPM SATWE 结构计算书", title_style))
    elements.append(Spacer(1, 10))

    # Insert structural images after title
    if images:
        elements.append(Paragraph("结构平面图", h1_style))
        for img_path in images:
            embed_path = _convert_bmp_to_png(img_path) if img_path.suffix.lower() == ".bmp" else img_path
            try:
                avail_w = doc.width
                img = RLImage(str(embed_path), width=avail_w, height=avail_w * 0.5)
                img.hAlign = "CENTER"
                elements.append(img)
                elements.append(Spacer(1, 8))
            except Exception:
                pass

    def _add_pdf_table(headers: List[str], rows: List[List[str]]) -> None:
        data = [headers] + rows
        col_count = len(headers)
        avail = doc.width
        col_w = [avail / col_count] * col_count
        t = Table(data, colWidths=col_w)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTNAME", (0, 0), (-1, -1), cn_font),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 8))

    # Section 1: Design Params
    params = out_data.get("wmass_params", {})
    if params:
        elements.append(Paragraph("一、设计参数总信息", h1_style))
        for key, label in [("design_params", "总体信息"), ("earthquake_params", "地震参数"), ("material_params", "材料参数")]:
            section = params.get(key)
            if section:
                elements.append(Paragraph(label, h2_style))
                for line in section.strip().split("\n"):
                    line = line.strip()
                    if line:
                        elements.append(Paragraph(line, body_style))

    # Section 2: Mass
    mass_table = out_data.get("mass_table", [])
    if mass_table:
        elements.append(Paragraph("二、各层质量", h1_style))
        _add_pdf_table(
            ["层号", "塔号", "恒载(t)", "活载(t)", "质量比"],
            [[str(r["floor"]), str(r["tower"]), str(r["dead_mass"]),
              str(r["live_mass"]), str(r["mass_ratio"])] for r in mass_table])

    # Section 3: Unit Mass
    unit_mass = out_data.get("unit_mass", [])
    if unit_mass:
        elements.append(Paragraph("三、单位面积质量分布", h1_style))
        _add_pdf_table(
            ["层号", "塔号", "单位质量(kg/m²)", "质量比"],
            [[str(r["floor"]), str(r["tower"]), str(r["unit_mass"]), str(r["mass_ratio"])]
             for r in unit_mass])

    # Section 4: Wind
    wind_load = out_data.get("wind_load", [])
    if wind_load:
        elements.append(Paragraph("四、风荷载信息", h1_style))
        _add_pdf_table(
            ["层号", "风X", "剪力X", "倾覆X", "风Y", "剪力Y", "倾覆Y"],
            [[str(r["floor"]), str(r["wind_x"]), str(r["shear_x"]), str(r["overturn_x"]),
              str(r["wind_y"]), str(r["shear_y"]), str(r["overturn_y"])] for r in wind_load])

    # Section 5: Modal
    wzq_periods = out_data.get("wzq_periods", [])
    modal_api = detailed.get("modal_analysis", [])
    if wzq_periods:
        elements.append(Paragraph("五、模态分析", h1_style))
        _add_pdf_table(
            ["振型号", "周期(s)", "方向角", "平动", "扭转"],
            [[str(m["mode"]), str(m["period"]), str(m["angle"]),
              str(m["translation"]), str(m["torsion"])] for m in wzq_periods])
    elif modal_api:
        elements.append(Paragraph("五、模态分析", h1_style))
        _add_pdf_table(
            ["振型号", "周期(s)", "方向角", "扭转系数", "X侧移", "Y侧移"],
            [[str(m["index"]), str(m["period_s"]), str(m["angle"]),
              str(m["torsion_ratio"]), str(m["x_side"]), str(m["y_side"])]
             for m in modal_api])

    # Section 6: CQC
    wzq_base_shear = out_data.get("wzq_base_shear", [])
    if wzq_base_shear:
        elements.append(Paragraph("六、地震楼层反应(CQC)", h1_style))
        for direction in ["X", "Y"]:
            dir_data = [r for r in wzq_base_shear if r["direction"] == direction]
            if dir_data:
                elements.append(Paragraph(f"{direction}方向", h2_style))
                _add_pdf_table(
                    ["层号", "塔号", "F(kN)", "V(kN)"],
                    [[str(r["floor"]), str(r["tower"]), str(r["F"]), str(r["V"])]
                     for r in dir_data])

    # Section 7: Shear Weight
    shear_weight = out_data.get("wzq_shear_weight_ratio", [])
    if shear_weight:
        elements.append(Paragraph("七、楼层剪重比", h1_style))
        _add_pdf_table(
            ["层号", "塔号", "X向", "Y向"],
            [[str(r["floor"]), str(r["tower"]), str(r["ratio_x"]), str(r["ratio_y"])]
             for r in shear_weight])

    # Section 8: Displacement
    wdisp_cases = out_data.get("wdisp_cases", [])
    if wdisp_cases:
        elements.append(Paragraph("八、层间位移", h1_style))
        for case in wdisp_cases:
            elements.append(Paragraph(f"工况{case['case_num']}: {case['title']}", h2_style))
            if case["floors"]:
                _add_pdf_table(
                    ["层号", "塔号", "最大位移", "平均位移"],
                    [[str(f["floor"]), str(f["tower"]), str(f["max_disp"]), str(f["ave_disp"])]
                     for f in case["floors"]])

    # Section 9: Stiffness
    stiff_api = detailed.get("story_stiffness", [])
    if stiff_api:
        elements.append(Paragraph("九、层刚度", h1_style))
        _add_pdf_table(
            ["层号", "RJX(kN/m)", "RJY(kN/m)", "比值X", "比值Y"],
            [[str(s["floor_index"]), str(s["RJX"]), str(s["RJY"]),
              str(s["ratio_x"]), str(s["ratio_y"])] for s in stiff_api])

    # Section 10: Overturning
    overturning = out_data.get("overturning", [])
    if overturning:
        elements.append(Paragraph("十、抗倾覆验算", h1_style))
        _add_pdf_table(
            ["工况", "抗倾覆力矩", "倾覆力矩", "比值", "零应力区(%)"],
            [[str(r["case"]), str(r["Mr"]), str(r["Mov"]), str(r["ratio"]), str(r["zero_stress"])]
             for r in overturning])

    # Section 11: Stability
    stability = out_data.get("stability", [])
    if stability:
        elements.append(Paragraph("十一、结构整体稳定验算", h1_style))
        _add_pdf_table(
            ["层号", "X刚度", "Y刚度", "层高", "X刚重比", "Y刚重比"],
            [[str(r["floor"]), str(r["stiff_x"]), str(r["stiff_y"]), str(r["height"]),
              str(r["ratio_x"]), str(r["ratio_y"])] for r in stability])

    # Section 12: Column Design
    wpj_columns = out_data.get("wpj_columns", [])
    if wpj_columns:
        elements.append(Paragraph("十四、柱配筋验算结果", h1_style))
        _add_pdf_table(
            ["编号", "截面", "轴压比", "配筋率(%)", "配箍率(%)"],
            [[str(c["id"]), f"{c.get('width', '')}*{c.get('height', '')}",
              str(c.get("axial_ratio", "")), str(c.get("reinforce_ratio", "")),
              str(c.get("hoop_ratio", ""))] for c in wpj_columns])

    # Section 13: Beam Design
    wpj_beams = out_data.get("wpj_beams", [])
    if wpj_beams:
        elements.append(Paragraph("十五、梁配筋验算结果", h1_style))
        _add_pdf_table(
            ["编号", "截面", "上部配筋", "下部配筋"],
            [[str(b["id"]), f"{b.get('width', '')}*{b.get('height', '')}",
              str(b.get("top_reinforce", "")), str(b.get("btm_reinforce", ""))]
             for b in wpj_beams])

    # Section: Exceedance
    wgcpj = out_data.get("wgcpj", {})
    exceed_items = wgcpj.get("items", [])
    if exceed_items:
        elements.append(Paragraph("十六、超限信息", h1_style))
        for item in exceed_items:
            elements.append(Paragraph(item, body_style))

    # Section: API Summary
    beams = detailed.get("beam_design", {})
    if beams.get("total_beams", 0) > 0:
        elements.append(Paragraph("十七、梁设计统计", h1_style))
        elements.append(Paragraph(f"总梁数: {beams['total_beams']}", body_style))
        elements.append(Paragraph(f"最大剪压比: {beams['max_shear_compression_ratio']}", body_style))

    cols = detailed.get("column_design", {})
    if cols.get("total_columns", 0) > 0:
        elements.append(Paragraph("十八、柱设计统计", h1_style))
        elements.append(Paragraph(f"总柱数: {cols['total_columns']}", body_style))
        elements.append(Paragraph(f"最大轴压比: {cols['max_axial_compression_ratio']}", body_style))

    doc.build(elements)
    return output_path


# ── PDF conversion fallback ──────────────────────────────────────────────


def _find_wps_exe() -> Optional[Path]:
    base = Path(os.environ.get("LOCALAPPDATA", "")) / "Kingsoft" / "WPS Office"
    if not base.is_dir():
        return None
    for ver_dir in sorted(base.iterdir(), reverse=True):
        candidate = ver_dir / "office6" / "wps.exe"
        if candidate.is_file():
            return candidate
    return None


def _convert_docx_to_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    wps_exe = _find_wps_exe()
    if wps_exe:
        try:
            proc = subprocess.Popen(
                [str(wps_exe), str(docx_path), "/ExportPDF", str(pdf_path)],
                stdout=subprocess.PIPE, stderr=subprocess.PIPE,
            )
            proc.wait(timeout=30)
            if pdf_path.is_file():
                return pdf_path
        except Exception:
            pass
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        wps = win32com.client.Dispatch("kwps.Application")
        wps.Visible = False
        doc = wps.Documents.Open(str(docx_path.resolve()))
        doc.ExportAsFixedFormat(str(pdf_path.resolve()), 0)
        doc.Close(False)
        wps.Quit()
        pythoncom.CoUninitialize()
        if pdf_path.is_file():
            return pdf_path
    except Exception:
        pass
    return None


# ── Main entry point ────────────────────────────────────────────────────


def run_analysis(model: Dict[str, Any], parameters: Dict[str, Any]) -> Dict[str, Any]:
    jws_path = _resolve_jws_path(model, parameters)
    warnings: List[str] = []
    project_dir = jws_path.parent

    # ── Read .OUT files ──────────────────────────────────────────────
    wmass_text = _read_out_file(project_dir, "WMASS.OUT")
    wzq_text = _read_out_file(project_dir, "WZQ.OUT")
    wdisp_text = _read_out_file(project_dir, "WDISP.OUT")
    wgcpj_text = _read_out_file(project_dir, "WGCPJ.OUT")
    wpj_data = _read_all_wpj(project_dir)

    out_file_data: Dict[str, Any] = {}

    if wmass_text:
        out_file_data["wmass_params"] = _parse_wmass_sections(wmass_text)
        out_file_data["mass_table"] = _parse_mass_table(wmass_text)
        out_file_data["total_mass"] = _parse_total_mass(wmass_text)
        out_file_data["member_counts"] = _parse_member_counts(wmass_text)
        out_file_data["wind_load"] = _parse_wind_load(wmass_text)
        out_file_data["floor_dimensions"] = _parse_floor_dimensions(wmass_text)
        out_file_data["unit_mass"] = _parse_unit_mass(wmass_text)
        out_file_data["stiffness_info"] = _parse_stiffness_info(wmass_text)
        out_file_data["overturning"] = _parse_overturning(wmass_text)
        out_file_data["stability"] = _parse_stability(wmass_text)
        out_file_data["stability_conclusion"] = _parse_stability_conclusion(wmass_text)
        out_file_data["shear_capacity"] = _parse_shear_capacity(wmass_text)
        out_file_data["comfort"] = _parse_comfort(wmass_text)

    if wzq_text:
        out_file_data["wzq_periods"] = _parse_wzq_periods(wzq_text)
        out_file_data["wzq_direction_factors"] = _parse_wzq_direction_factors(wzq_text)
        out_file_data["wzq_base_shear"] = _parse_wzq_base_shear(wzq_text)
        out_file_data["wzq_effective_mass"] = _parse_wzq_effective_mass(wzq_text)
        out_file_data["wzq_min_shear_ratio"] = _parse_wzq_min_shear_ratio(wzq_text)
        out_file_data["wzq_shear_weight_ratio"] = _parse_wzq_shear_weight_ratio(wzq_text)

    if wdisp_text:
        out_file_data["wdisp_cases"] = _parse_wdisp_cases(wdisp_text)

    if wgcpj_text:
        out_file_data["wgcpj"] = _parse_wgcpj(wgcpj_text)

    all_wpj_columns: List[Dict[str, Any]] = []
    all_wpj_beams: List[Dict[str, Any]] = []
    for wpj in wpj_data:
        all_wpj_columns.extend(_parse_wpj_columns(wpj["content"]))
        all_wpj_beams.extend(_parse_wpj_beams(wpj["content"]))
    if all_wpj_columns:
        out_file_data["wpj_columns"] = all_wpj_columns
    if all_wpj_beams:
        out_file_data["wpj_beams"] = all_wpj_beams

    # ── Find structural images ───────────────────────────────────────
    project_images = _find_project_images(project_dir)

    # ── Try APIPyInterface (supplementary) ───────────────────────────
    api_data: Dict[str, Any] = {}
    try:
        import APIPyInterface
        result = APIPyInterface.ResultData()
        ret = result.InitialResult(str(jws_path))
        if ret != 0:
            warnings.append(f"InitialResult returned non-zero: {ret}")

        api_data["modal_analysis"] = _extract_modal(result)
        api_data["story_stiffness"] = _extract_story_stiffness(result)
        api_data["story_drift"] = _extract_story_drift(result)
        api_data["base_shear"] = _extract_base_shear(result)
        api_data["story_mass"] = _extract_story_mass(result)
        api_data["stiff_weight_ratio"] = _extract_stiff_weight_ratio(result)
        api_data["beam_design"] = _extract_beam_design(result)
        api_data["column_design"] = _extract_column_design(result)
        result.ClearResult()
    except ImportError:
        warnings.append("APIPyInterface not available, using .OUT file data only")
        for key in ["modal_analysis", "story_stiffness", "story_drift", "base_shear",
                     "story_mass", "stiff_weight_ratio", "beam_design", "column_design"]:
            api_data.setdefault(key, [])
        api_data.setdefault("story_drift", {"earthquake": {}, "wind": {}, "limit_value": None})
        api_data.setdefault("base_shear", {"entries": [], "shear_weight_limit": None})
        api_data.setdefault("stiff_weight_ratio", {"entries": [], "limit_value": None})
    except Exception as exc:
        warnings.append(f"APIPyInterface error: {exc}")
        for key in ["modal_analysis", "story_stiffness", "story_drift", "base_shear",
                     "story_mass", "stiff_weight_ratio", "beam_design", "column_design"]:
            api_data.setdefault(key, [])

    # ── Build report ─────────────────────────────────────────────────
    detailed = {
        "modal_analysis": api_data.get("modal_analysis", []),
        "story_stiffness": api_data.get("story_stiffness", []),
        "story_drift": api_data.get("story_drift", {}),
        "base_shear": api_data.get("base_shear", {}),
        "story_mass": api_data.get("story_mass", []),
        "stiff_weight_ratio": api_data.get("stiff_weight_ratio", {}),
        "beam_design": api_data.get("beam_design", {}),
        "column_design": api_data.get("column_design", {}),
        "code_exceedance": [],
        "out_file_data": out_file_data,
        "images": [str(p) for p in project_images],
    }

    report: Dict[str, Any] = {
        "status": "success",
        "summary": {
            "engine": "pkpm-calcbook",
            "jws_path": str(jws_path),
            "mode_count": len(api_data.get("modal_analysis", [])),
            "beam_total": api_data.get("beam_design", {}).get("total_beams", 0),
            "column_total": api_data.get("column_design", {}).get("total_columns", 0),
            "out_files_parsed": len([k for k in out_file_data if out_file_data[k]]),
            "wpj_column_count": len(all_wpj_columns),
            "wpj_beam_count": len(all_wpj_beams),
        },
        "detailed": detailed,
        "warnings": warnings,
    }

    report["markdown"] = _generate_markdown(report)

    # Generate Word document
    output_dir = Path(parameters.get("output_dir", "")).resolve() or jws_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / f"{jws_path.stem}_计算书.docx"

    try:
        _generate_docx(report, docx_path, images=project_images)
        report["summary"]["docx_path"] = str(docx_path)
    except Exception as exc:
        warnings.append(f"Word generation failed: {exc}")

    # Generate PDF
    pdf_path = output_dir / f"{jws_path.stem}_计算书.pdf"
    try:
        _generate_pdf(report, pdf_path, images=project_images)
        report["summary"]["pdf_path"] = str(pdf_path)
    except ImportError:
        try:
            result_path = _convert_docx_to_pdf(docx_path)
            if result_path:
                report["summary"]["pdf_path"] = str(result_path)
            else:
                warnings.append("PDF generation skipped: reportlab not installed and WPS conversion failed")
        except Exception as exc:
            warnings.append(f"PDF generation failed: {exc}")
    except Exception as exc:
        warnings.append(f"PDF generation failed: {exc}")

    return report
